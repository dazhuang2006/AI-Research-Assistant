"""
统一文档解析器

把 PDF、DOCX、HTML、TXT 统一转成纯文本，再切成带重叠的文本分块。
"""
import os
from typing import List, Tuple

from bs4 import BeautifulSoup
from docx import Document
import pdfplumber


def extract_text_from_pdf(file_path: str) -> str:
    """
    从 PDF 文件中逐页提取文本

    Args:
        file_path: PDF 文件路径

    Returns:
        提取出的文本内容
    """
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    return text


def extract_text_from_docx(file_path: str) -> str:
    """
    从 DOCX 文件中提取段落与表格文本

    Args:
        file_path: DOCX 文件路径

    Returns:
        提取出的文本内容
    """
    doc = Document(file_path)
    parts = []

    # 提取段落文本
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    # 提取表格文本
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_text_from_html(file_path: str) -> str:
    """
    从 HTML 文件中提取纯文本，去掉脚本与样式

    Args:
        file_path: HTML 文件路径

    Returns:
        提取出的文本内容
    """
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        html_content = f.read()

    soup = BeautifulSoup(html_content, "html.parser")
    for tag in soup(["script", "style", "meta", "link"]):
        tag.decompose()

    return soup.get_text(separator="\n", strip=True)


def extract_text_from_txt(file_path: str) -> str:
    """
    从 TXT 文件中直接读取文本

    Args:
        file_path: TXT 文件路径

    Returns:
        文本内容
    """
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text_from_file(file_path: str) -> Tuple[str, str]:
    """
    根据扩展名自动选择解析器

    Args:
        file_path: 文件路径

    Returns:
        (提取出的文本, 文件类型) 元组

    Raises:
        ValueError: 文件格式不支持或没有提取到有效文本
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    parsers = {
        ".pdf": (extract_text_from_pdf, "PDF"),
        ".docx": (extract_text_from_docx, "DOCX"),
        ".html": (extract_text_from_html, "HTML"),
        ".htm": (extract_text_from_html, "HTML"),
        ".txt": (extract_text_from_txt, "TXT"),
    }

    if ext not in parsers:
        raise ValueError(f"不支持的文件格式: {ext}")

    parser_func, file_type = parsers[ext]
    text = parser_func(file_path)

    if not text or len(text.strip()) < 10:
        raise ValueError(f"未能从 {file_type} 文件中提取到有效文本")

    return text, file_type


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """
    将长文本切成带重叠的分块

    重叠的目的是避免一句话在分块边界被切断，从而丢失语义。

    Args:
        text: 原始文本
        chunk_size: 每个分块的字符数
        overlap: 相邻分块重叠的字符数

    Returns:
        文本分块列表
    """
    if not text:
        return []

    chunks = []
    start = 0
    text_length = len(text)

    while start < text_length:
        chunk = text[start:start + chunk_size]
        if chunk.strip():
            chunks.append(chunk)
        start += chunk_size - overlap

    return chunks


# 受支持的文件扩展名
SUPPORTED_EXTENSIONS = [".pdf", ".docx", ".html", ".htm", ".txt"]
